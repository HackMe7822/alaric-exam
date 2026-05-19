from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Theme colours ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x00, 0x2B, 0x5C)
MID_BLUE   = RGBColor(0x00, 0x78, 0xD4)
ACCENT     = RGBColor(0x00, 0xBC, 0xF2)
GREEN      = RGBColor(0x00, 0xA8, 0x6B)
RED        = RGBColor(0xE7, 0x48, 0x56)
AMBER      = RGBColor(0xF5, 0x9E, 0x0B)
LIGHT_GREY = RGBColor(0xF0, 0xF4, 0xF8)
MID_GREY   = RGBColor(0x64, 0x78, 0x8B)
TEXT_DARK  = RGBColor(0x1A, 0x20, 0x2C)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: heading styles ────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    shade_paragraph(p, '002B5C')
    run = p.add_run(f'  {text}')
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = MID_BLUE
    run.font.name = 'Calibri'
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '0078D4')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_DARK
    run.font.name = 'Calibri'
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_DARK
    run.font.name = 'Calibri'
    return p

def add_table(headers, rows, col_widths=None, header_color='0078D4'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name  = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            cell.width = Inches(col_widths[i])
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        fill = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size  = Pt(9)
            run.font.color.rgb = TEXT_DARK
            run.font.name  = 'Calibri'
            if col_widths:
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()
    return table

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'D0D8E0')
    pBdr.append(bot)
    pPr.append(pBdr)

def info_box(label, text, color='0078D4'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    shade_paragraph(p, 'E8F4FD')
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size  = Pt(10)
    r1.font.color.rgb = MID_BLUE
    r1.font.name  = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = TEXT_DARK
    r2.font.name  = 'Calibri'

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '002B5C')
r = p.add_run('  ')
r.font.size = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
shade_paragraph(p, '002B5C')
r = p.add_run('ALARIC EXAM PLATFORM')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = WHITE
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
shade_paragraph(p, '002B5C')
r = p.add_run('Enterprise Examination Management System')
r.font.size = Pt(16)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
shade_paragraph(p, '002B5C')
r = p.add_run('  ')
r.font.size = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
shade_paragraph(p, '002B5C')
r = p.add_run('─' * 52)
r.font.color.rgb = MID_BLUE
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
shade_paragraph(p, '002B5C')
r = p.add_run('  ')
r.font.size = Pt(8)

for line, sz, col in [
    ('Project Specification Document', 13, ACCENT),
    ('Version 1.0  |  CONFIDENTIAL', 10, MID_GREY),
    (f'Date: {datetime.date.today().strftime("%B %d, %Y")}', 10, MID_GREY),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    shade_paragraph(p, '002B5C')
    r = p.add_run(line)
    r.font.size = Pt(sz)
    r.font.color.rgb = col
    r.font.name = 'Calibri'

for _ in range(8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    shade_paragraph(p, '002B5C')
    r = p.add_run(' ')
    r.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Executive Summary')
body(
    'Alaric Exam is a full-featured, enterprise-grade examination management platform '
    'designed to create, deliver, proctor, and evaluate technical and professional assessments. '
    'Built for organisations that need reliable, scalable, and secure examination infrastructure, '
    'Alaric Exam replaces fragmented exam tools with a single, cohesive platform covering '
    'the entire lifecycle — from authoring questions to delivering results and certificates.'
)
body(
    'The platform supports eight question types, multi-role administration, automated scoring, '
    'manual review workflows, AI and multi-monitor proctoring, Microsoft Graph email integration, '
    'a self-service candidate portal, advanced analytics, REST API access, and gamification — '
    'all deployable locally for testing and migrating to Cloudflare\'s edge infrastructure for production.'
)

info_box('Platform Name',    'Alaric Exam')
info_box('Version',          '1.0 — Local (Node.js + SQLite)')
info_box('Production Target','Cloudflare Workers + D1 + R2')
info_box('Email Service',    'Microsoft Graph API (Azure App Registration)')
info_box('Document Status',  'Pre-Build Specification — Approved for Development')

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Project Overview')

h2('2.1 Purpose')
body(
    'Alaric Exam addresses a common organisational challenge: conducting controlled, '
    'credible technical assessments for engineers, candidates, and employees without '
    'relying on third-party SaaS platforms that lack customisation, data ownership, '
    'or integration flexibility.'
)

h2('2.2 Core Objectives')
for obj in [
    'Provide administrators with a comprehensive exam authoring and management interface.',
    'Deliver a professional, secure, and accessible exam experience to candidates.',
    'Automate scoring for objective questions while supporting manual review for subjective answers.',
    'Prevent academic dishonesty through browser-level proctoring, AI-use detection, and multi-monitor controls.',
    'Integrate natively with Microsoft Graph API for automated email delivery of invitations and results.',
    'Expose a REST API for integration with HR systems, LMS platforms, and automation tools.',
    'Support a self-service candidate portal for result history, certificate download, and retake requests.',
    'Provide actionable analytics for exam managers to continuously improve question quality and exam design.',
]:
    bullet(obj)

h2('2.3 Target Users')
add_table(
    ['Role', 'Description', 'Access Level'],
    [
        ['Super Admin',    'Full platform control — users, settings, email config, all data',         'Unrestricted'],
        ['Exam Manager',   'Create/publish exams, manage candidates, view results',                   'Exam + Candidate scope'],
        ['Checker',        'Review text/essay answers, assign marks and remarks',                     'Assigned submissions only'],
        ['Viewer',         'Read-only access to results and analytics — no editing',                  'Read-only'],
        ['Candidate',      'Take assigned exams, view own results via portal',                        'Own data only'],
    ],
    col_widths=[1.4, 3.8, 1.6]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Technology Stack')

h2('3.1 Local Development Stack')
add_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Runtime',        'Node.js 18+',                  'Server-side JavaScript runtime'],
        ['Framework',      'Express.js',                   'HTTP routing, middleware, API'],
        ['Database',       'SQLite via better-sqlite3',    'Local relational data storage (22 tables)'],
        ['Authentication', 'JWT (HttpOnly cookies)',        'Stateless session management'],
        ['2FA',            'TOTP (speakeasy library)',      'Time-based one-time passwords for admin login'],
        ['PDF Generation', 'jsPDF + jsPDF-AutoTable',      'Client-side result PDF and certificate generation'],
        ['Email',          'Microsoft Graph REST API',      'Transactional email via Azure app registration'],
        ['File Storage',   'Local /uploads directory',     'CSV uploads, webcam snapshots, attachments'],
        ['Frontend',       'Vanilla HTML / CSS / JS',       'Zero build tools — fast, portable, inspectable'],
        ['Service Worker', 'Browser Service Worker API',   'Offline PWA support, answer caching'],
    ],
    col_widths=[1.5, 2.2, 3.1]
)

h2('3.2 Production Migration Path (Cloudflare)')
add_table(
    ['Local',              'Cloudflare Equivalent',    'Change Required'],
    [
        ['Express.js',     'Cloudflare Workers',       'Minimal — route handlers map 1:1'],
        ['SQLite file',    'Cloudflare D1',             'None — identical SQL syntax'],
        ['/uploads folder','Cloudflare R2',             'Replace fs.writeFile with R2 put()'],
        ['JWT cookies',    'Same',                      'No change'],
        ['Graph API',      'Same',                      'No change'],
        ['Static assets',  'Cloudflare Pages',          'Copy public/ folder'],
    ],
    col_widths=[2.0, 2.2, 2.6]
)

h2('3.3 Project Structure')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
shade_paragraph(p, 'F0F4F8')
run = p.add_run(
    'alaric-exam/\n'
    '├── server.js                   ← Express entry point\n'
    '├── package.json\n'
    '├── .env                        ← Secrets (Graph API, JWT)\n'
    '├── database/\n'
    '│   ├── schema.sql              ← 22-table schema\n'
    '│   └── alaric.db               ← Auto-generated SQLite\n'
    '├── src/\n'
    '│   ├── routes/admin/           ← All admin API endpoints\n'
    '│   ├── routes/public/          ← Catalog + exam-taking\n'
    '│   ├── middleware/             ← Auth, roles, rate-limit\n'
    '│   └── services/              ← Email, PDF, scorer, token\n'
    '└── public/\n'
    '    ├── admin/                  ← ~20 admin HTML pages\n'
    '    ├── exam/                   ← Candidate exam page\n'
    '    ├── portal/                 ← Candidate self-service\n'
    '    ├── catalog/                ← Public exam listing\n'
    '    └── assets/css|js/          ← Shared styles + scripts'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = TEXT_DARK

# ══════════════════════════════════════════════════════════════════════════════
# 4. QUESTION TYPES
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Question Types')
body('Alaric Exam supports eight distinct question formats to assess a broad range of competencies:')

add_table(
    ['#', 'Type', 'Scoring', 'Checker Required', 'Use Case'],
    [
        ['1', 'Multiple Choice (Single)',   'Automatic',       'No',  'Knowledge recall, conceptual understanding'],
        ['2', 'Multiple Select (Multi)',    'Automatic',       'No',  'Complex scenarios with multiple valid answers'],
        ['3', 'True / False',               'Automatic',       'No',  'Quick factual verification'],
        ['4', 'Fill in the Blank',          'Automatic',       'No',  'Terminology, syntax, command recall'],
        ['5', 'Text / Essay',               'Manual',          'Yes', 'Explanatory, analytical, open-ended responses'],
        ['6', 'Drag & Drop Ordering',       'Automatic',       'No',  'Sequence, process, priority ordering'],
        ['7', 'Match the Pairs',            'Automatic',       'No',  'Association, definition matching'],
        ['8', 'File Upload',                'Manual',          'Yes', 'Diagrams, code files, written documents'],
    ],
    col_widths=[0.3, 2.0, 1.2, 1.4, 2.0]
)

body('Each question supports:')
for f in [
    'Optional code block display (monospace, syntax-highlighted) above the question text',
    'Difficulty tag: Easy / Medium / Hard (used for auto-exam generation and analytics)',
    'Maximum marks (default 1 for auto-scored; configurable for manual types)',
    'Explanation field (shown to candidate after exam review)',
    'Per-question time override when exam is configured in per-question time mode',
    'Section/domain assignment for domain-level scoring breakdown',
    'Internal checker reference answer (private, never shown to candidate)',
]:
    bullet(f)

# ══════════════════════════════════════════════════════════════════════════════
# 5. EXAM CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Exam Configuration')
body('Every exam is independently configurable across the following dimensions:')

h2('5.1 General Settings')
add_table(
    ['Setting', 'Options / Description'],
    [
        ['Title & Description',     'Full text + candidate-facing instructions'],
        ['Mode',                    'Practice (open catalog) / Assessed (private link only)'],
        ['Status',                  'Draft / Published / Archived'],
        ['Availability Window',     'Optional open date + close date'],
        ['Max Attempts',            '1, 3, 5, or Unlimited (for practice mode)'],
        ['Passing Score',           'Percentage threshold (default 70%)'],
        ['Show Results Immediately','Yes = instant results; No = held until checker reviews text answers'],
        ['Certificate on Pass',     'Toggle — generates branded PDF certificate'],
        ['Exam Password',           'Optional extra access code required beyond the unique link'],
    ],
    col_widths=[2.2, 4.6]
)

h2('5.2 Timing')
add_table(
    ['Setting', 'Description'],
    [
        ['Time Mode',              'Total exam countdown OR per-question countdown'],
        ['Total Time Limit',       'Hours and minutes (e.g. 3:00:00 for AZ-104 style)'],
        ['Per-Question Time',      'Seconds per question; auto-advances on expiry'],
        ['Per-Question Override',  'Individual questions can override the global per-question time'],
    ],
    col_widths=[2.2, 4.6]
)

h2('5.3 Scoring')
add_table(
    ['Setting', 'Description'],
    [
        ['Negative Marking',       'Deduction per wrong answer (configurable: e.g. -0.25)'],
        ['Weighted Sections',      'Assign percentage weight per section (must total 100%)'],
        ['Bonus Questions',        'Optional extra-credit questions — add to score, not to denominator'],
        ['Grace Marks',            'Admin can add bonus marks to a completed submission post-review'],
        ['Partial Credit (Multi)', 'Full marks only / proportional (1 mark per correct option selected)'],
    ],
    col_widths=[2.2, 4.6]
)

h2('5.4 Question Delivery')
add_table(
    ['Setting', 'Description'],
    [
        ['Question Order',         'Fixed sequence OR randomised per candidate'],
        ['Option Order',           'Fixed OR shuffled per candidate'],
        ['Question Pools',         'Draw N random questions from a larger pool per section'],
        ['Auto-Generate from Bank','Select count per difficulty level; system builds exam automatically'],
    ],
    col_widths=[2.2, 4.6]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. ANTI-CHEAT & PROCTORING
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Anti-Cheat & Proctoring')

h2('6.1 Browser-Level Controls')
add_table(
    ['Control', 'Implementation', 'Violation Action'],
    [
        ['Right-click disabled',       'contextmenu event prevented',                       'Silent — no action'],
        ['Copy / Paste blocked',       'copy, paste, cut events prevented on exam area',    'Silent log'],
        ['Text selection disabled',    'CSS user-select: none on question area',             'Silent'],
        ['Keyboard shortcut blocking', 'F12, Ctrl+Shift+I, Ctrl+U, Ctrl+C, Ctrl+V, Ctrl+A','Silent log'],
        ['DevTools detection',         'Window size delta + console timing trick',           'Warning overlay + log'],
        ['Print / Screenshot deterrent','CSS @media print hidden + overlay on blur',        'Warning'],
    ],
    col_widths=[2.1, 2.5, 2.2]
)

h2('6.2 Focus & Window Monitoring')
add_table(
    ['Detection', 'Method', 'Threshold & Action'],
    [
        ['Tab switch / App switch', 'document.visibilityState + window.blur events',         '3 violations → auto-submit warning'],
        ['Fullscreen enforcement',  'Fullscreen API — exam requires and maintains fullscreen','Exit = instant pause; 3 exits → auto-submit'],
        ['Mouse leaving window',    'mouseleave on document body',                            'Logged with timestamp'],
        ['Window blur event',       'window.blur listener with timestamp',                   'Cumulative count shown to admin'],
    ],
    col_widths=[2.0, 2.6, 2.2]
)

h2('6.3 Multi-Monitor Detection')
add_table(
    ['Detection', 'Method', 'Action'],
    [
        ['Multiple displays present',  'screen.isExtended (Chrome 100+) at exam start',          'Admin-configurable: warn-only OR block exam'],
        ['Window on second monitor',   'window.screenLeft / screenX outside primary bounds',     'Pause exam + alert every 10 seconds'],
        ['Screen extended permission', 'Window Management API — request explicit permission',    'Decline = treat as multi-monitor warning'],
    ],
    col_widths=[2.2, 2.8, 1.8]
)

h2('6.4 AI Use Detection')
add_table(
    ['Detection', 'Method', 'Flag in Result'],
    [
        ['Paste into text answer',     'input event: sudden large text block (>50 chars, 0 keystrokes)', 'Yes — checker sees "Possible AI paste" tag'],
        ['Typing rhythm anomaly',      'Keypress rate analysis — human avg 40–80 WPM baseline',          'Yes — flagged for checker review'],
        ['Clipboard content monitor',  'paste event captures pasted text length + timing',               'Logged with character count'],
        ['Answer velocity scoring',    'Time from question displayed to answer submitted',                'Statistical outlier flagged'],
    ],
    col_widths=[2.2, 2.8, 1.8]
)

h2('6.5 Webcam Proctoring')
for b in [
    'Periodic snapshots taken via getUserMedia() at configurable intervals (e.g. every 2 minutes)',
    'Snapshots stored server-side, visible in admin results panel per candidate',
    'Candidate shown consent notice before exam starts — exam cannot begin without granting permission',
    'Admin-configurable: Required / Optional / Disabled per exam',
    'Snapshot count and timestamps logged in Integrity Report',
]:
    bullet(b)

h2('6.6 Integrity Report')
body(
    'Every submission includes a machine-generated Integrity Report visible to admin and checkers, '
    'summarising all detected violation events with timestamps, counts, and severity ratings. '
    'The report is appended to the result PDF and stored in the database.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. CANDIDATE LINK SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Candidate Link System')

h2('7.1 Link Properties')
add_table(
    ['Property', 'Detail'],
    [
        ['Token format',     '32-character cryptographically random string (UUID v4 variant)'],
        ['URL format',       'http://localhost:3000/exam/{token}'],
        ['Usage',            'Single-use only — expires immediately on submission'],
        ['Expiry',           'Configurable per link (default 7 days from generation)'],
        ['Statuses',         'Pending / In Progress / Completed / Expired / Revoked'],
        ['Re-open behaviour','Completed link shows: "This exam has already been submitted"'],
        ['Expired behaviour','Shows: "This link has expired. Contact your administrator"'],
    ],
    col_widths=[2.0, 4.8]
)

h2('7.2 Link Management (Admin)')
for b in [
    'Single invite: select candidate + exam → token generated → copy link or send email immediately',
    'Bulk invite: upload CSV (name, email columns) → select exam → system generates all tokens → sends batch invitation emails',
    'Per-link actions: Resend invitation email, Extend expiry date, Revoke link, Copy link to clipboard',
    'Duplicate detection: warns if same email already has an active link for the same exam',
    'Link history: full log of all links generated, sent, and their current status',
]:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 8. ADMIN PANEL
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Admin Panel')
body('The admin panel is accessible at /admin and is secured by role-based JWT authentication with optional 2FA.')

h2('8.1 Page Inventory')
add_table(
    ['Page', 'Path', 'Roles', 'Description'],
    [
        ['Login',            '/admin/login',       'All',         'Email + password + optional TOTP 2FA'],
        ['Dashboard',        '/admin',             'All',         'KPI cards, pending reviews badge, activity feed'],
        ['Exams List',       '/admin/exams',       'Mgr+',        'All exams with status, actions, duplicate'],
        ['Exam Editor',      '/admin/exams/:id',   'Mgr+',        'Settings, sections, questions tabs'],
        ['Question Bank',    '/admin/bank',        'Mgr+',        'Global reusable questions library'],
        ['Candidates',       '/admin/candidates',  'Mgr+',        'Candidate records, links, bulk invite'],
        ['Results',          '/admin/results',     'Viewer+',     'All submissions, filter, export, re-download'],
        ['Review Queue',     '/admin/review',      'Checker+',    'Text and file answer review interface'],
        ['Analytics',        '/admin/analytics',   'Viewer+',     'Charts, item analysis, cohort compare'],
        ['Users',            '/admin/users',       'Super Admin', 'Admin accounts, roles, sessions'],
        ['Email Config',     '/admin/email',       'Super Admin', 'Graph credentials, templates, test send'],
        ['Settings',         '/admin/settings',    'Super Admin', 'Branding, certificate, defaults, retention'],
        ['Audit Log',        '/admin/audit',       'Super Admin', 'Full action log across all users'],
        ['API Keys',         '/admin/api',         'Super Admin', 'Generate and manage REST API keys'],
    ],
    col_widths=[1.6, 1.8, 1.4, 2.7]
)

h2('8.2 Exam Editor — Tabs')
h3('Settings Tab')
body('General metadata, mode, timing, scoring, anti-cheat toggles, certificate, availability window.')
h3('Sections Tab')
body('Add, rename, reorder, colour-code, and delete exam sections. Each section can have a weighting percentage.')
h3('Questions Tab')
for b in [
    'Inline question cards with drag-and-drop reorder',
    'Add question button opens type selector then full question form',
    'Bulk import via CSV with per-row validation preview',
    'Download sample CSV button always visible',
    'Add from Question Bank — search and copy questions from global bank',
    'Question review workflow toggle: Draft / Peer Review / Approved status per question',
    'Version history: see previous versions of any question with rollback',
]:
    bullet(b)

h2('8.3 Review Queue (Checker Interface)')
body(
    'The review queue lists all submissions containing unreviewed text or file answers. '
    'Checkers see only submissions assigned to their scope. For each answer, the checker:'
)
for b in [
    'Reads the candidate\'s response alongside the private sample/reference answer',
    'Selects verdict: Correct / Partially Correct / Incorrect',
    'For Partially Correct: enters marks (0 to question\'s maximum) — mandatory',
    'For Partially Correct and Incorrect: enters remarks — mandatory',
    'Submits review — when all questions in submission are reviewed, final score auto-calculates',
    'If two checkers disagree (multi-checker mode), submission escalates to Super Admin',
]:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 9. EXAM TAKING EXPERIENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Exam Taking Experience')

h2('9.1 Candidate Flow')
add_table(
    ['Step', 'Description'],
    [
        ['1. Link opens',        'Token validated — candidate name pre-populated from invite record'],
        ['2. Instructions',      'Exam title, duration, question count, section breakdown, anti-cheat notice'],
        ['3. Fullscreen prompt', 'Browser requests fullscreen — exam cannot start without it'],
        ['4. Webcam consent',    'If webcam proctoring enabled — camera permission requested'],
        ['5. Exam begins',       'Timer starts, first question displayed, navigation panel available'],
        ['6. Navigation',        'Left panel: question grid by section, colour-coded status (answered/flagged/current)'],
        ['7. Flag questions',    'Star flag for review; flagged shown in submit confirmation modal'],
        ['8. Submit',            'Confirmation modal: shows answered count, unanswered warning, flagged list'],
        ['9. Post-submit',       'If instant results: full score + review screen + PDF download; else: holding screen'],
        ['10. Email',            'Result PDF emailed automatically via Microsoft Graph'],
    ],
    col_widths=[1.6, 5.2]
)

h2('9.2 Question Rendering')
for b in [
    'Multiple Choice: radio button cards — single selection, clear visual selection state',
    'Multiple Select: checkbox cards — "Select all that apply" label shown',
    'True / False: two large button cards — True / False',
    'Fill in the Blank: sentence with input field(s) inline',
    'Text / Essay: full textarea, spell-check enabled, paste blocked if anti-cheat on',
    'Drag & Drop: draggable item cards in a ranked list',
    'Match the Pairs: two columns with connect/select interaction',
    'File Upload: drag-and-drop upload zone with file type restriction',
    'Code Block: syntax-highlighted monospace panel displayed above question text',
]:
    bullet(b)

h2('9.3 Accessibility & Quality of Life')
for b in [
    'Font size controls (+/-) in exam toolbar',
    'High contrast mode toggle',
    'Calculator widget (toggleable per exam)',
    'Save & resume: answers cached every 30 seconds — reconnection restores state',
    'Offline PWA: service worker caches exam, answers persist if connection drops briefly',
    'Progress bar: visual completion indicator',
    'Time warnings: yellow at 30 min remaining, red at 10 min',
]:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 10. CANDIDATE PORTAL
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Candidate Portal')
body(
    'Accessible at /portal — candidates log in with their email address and a one-time OTP '
    'sent to their inbox (no separate password to manage).'
)
add_table(
    ['Feature', 'Description'],
    [
        ['Exam history',          'All exams taken: date, score, pass/fail, time taken'],
        ['Result detail',         'Full question review with correct answers, explanations, checker remarks'],
        ['Certificate download',  'Download pass certificate PDF for any passed exam'],
        ['Retake request',        'Submit request to admin for another attempt; admin approves/denies'],
        ['Pre-exam materials',    'Study resources attached to upcoming exams visible before start'],
        ['Post-exam resources',   'Topic-specific learning links shown based on wrong-answer domains'],
        ['Learning paths',        'Sequence of exams — progress tracker, locked exams unlock on pass'],
        ['Leaderboard (practice)','Opt-in ranking on practice exam leaderboards'],
        ['Badges & streaks',      'Earned badges shown on profile; consecutive exam streak counter'],
    ],
    col_widths=[2.0, 4.8]
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. EMAIL SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Email System (Microsoft Graph API)')

h2('11.1 Configuration')
body(
    'Email is configured once in Admin → Email Config. Azure App Registration credentials '
    'are stored encrypted in the database.'
)
add_table(
    ['Field', 'Description'],
    [
        ['Tenant ID',      'Azure Active Directory tenant identifier'],
        ['Client ID',      'App registration client ID'],
        ['Client Secret',  'App registration secret (encrypted at rest)'],
        ['From Email',     'Sender address (e.g. exams@company.com)'],
        ['From Name',      'Display name (e.g. Alaric Exam Platform)'],
        ['Default CC',     'Comma-separated CC addresses for all result emails'],
        ['Test Send',      'Sends a test email immediately to verify credentials'],
    ],
    col_widths=[1.8, 5.0]
)

h2('11.2 Email Templates')
add_table(
    ['Template', 'Trigger', 'Key Variables'],
    [
        ['Exam Invitation',       'Admin sends link (single or bulk)',       '{{candidate_name}}, {{exam_title}}, {{exam_link}}, {{expires_date}}'],
        ['Link Expiry Reminder',  'X hours before link expires (configurable)','{{candidate_name}}, {{exam_title}}, {{hours_remaining}}, {{exam_link}}'],
        ['Result — Pass',         'Exam submitted + scoring complete',       '{{score}}, {{percentage}}, {{exam_title}}, PDF attached'],
        ['Result — Fail',         'Exam submitted + scoring complete',       '{{score}}, {{percentage}}, {{exam_title}}, PDF attached'],
        ['Checker Notification',  'Text answers pending review',             '{{exam_title}}, {{pending_count}}, {{review_link}}'],
        ['Retake Approved',       'Admin approves retake request',           '{{candidate_name}}, {{exam_title}}, {{new_exam_link}}'],
        ['Retake Denied',         'Admin denies retake request',             '{{candidate_name}}, {{exam_title}}, {{reason}}'],
    ],
    col_widths=[1.8, 2.2, 2.8]
)
body('All templates are editable HTML in the admin panel with live preview.')

# ══════════════════════════════════════════════════════════════════════════════
# 12. ANALYTICS & REPORTING
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Analytics & Reporting')

h2('12.1 Dashboard Metrics')
for b in [
    'Total exams published / total submissions / total candidates',
    'Pass rate this month vs last month (trend arrow)',
    'Pending review queue count (badge)',
    'Average completion time across all exams',
    'Recent activity feed (last 10 events)',
]:
    bullet(b)

h2('12.2 Per-Exam Analytics')
add_table(
    ['Report', 'Description'],
    [
        ['Score distribution',     'Histogram of score bands (0-10%, 10-20%, etc.)'],
        ['Pass rate trend',         'Monthly pass rate line chart'],
        ['Avg time per question',   'Bar chart — identify questions that take too long'],
        ['Item analysis',           'Per question: % correct, avg marks, discrimination index'],
        ['Difficulty index',        'Auto-calculated; questions auto-re-tagged if data diverges from manual tag'],
        ['Drop-off analysis',       'At which question do candidates abandon the exam'],
        ['AI flag frequency',       'Questions most frequently flagged for AI use detection'],
        ['Cohort comparison',       'Compare two batches on same exam side-by-side'],
    ],
    col_widths=[2.2, 4.6]
)

h2('12.3 Custom Report Builder')
body(
    'Drag-and-drop column selector: choose fields (candidate, exam, score, domain scores, time taken, '
    'violations, etc.), apply filters, sort, group by — save as named template. '
    'Export to Excel (.xlsx) or CSV. Schedule report to be emailed weekly.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 13. SECURITY
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Security')

add_table(
    ['Security Control', 'Implementation'],
    [
        ['Authentication',          'bcrypt password hashing (cost factor 12) + JWT HttpOnly cookies'],
        ['2FA',                     'TOTP via authenticator app (Google Authenticator compatible)'],
        ['SSO / Azure AD',          'OAuth2 login with Microsoft — admin can sign in with org account'],
        ['Role enforcement',        'Server-side middleware on every endpoint — no client-side-only guards'],
        ['Encrypted secrets',       'Graph API client secret + sensitive config encrypted in DB at rest'],
        ['Encrypted text answers',  'Candidate essay answers encrypted in DB (AES-256)'],
        ['Session management',      'Admin can view + force-terminate all active sessions'],
        ['Rate limiting',           'Login: 5 attempts / 15 minutes; API: configurable per key'],
        ['CSRF protection',         'SameSite=Strict cookies + CSRF tokens on state-changing forms'],
        ['Data retention',          'Auto-delete submissions older than configured months (GDPR)'],
        ['GDPR export',             'One-click export of all data for a specific candidate'],
        ['IP restriction',          'Optional per-exam: whitelist IP ranges allowed to access exam'],
        ['Audit log',               'Every admin action logged with user, timestamp, entity, before/after'],
        ['API key scoping',         'Each API key granted only specified permissions (read/write/admin)'],
    ],
    col_widths=[2.2, 4.6]
)

# ══════════════════════════════════════════════════════════════════════════════
# 14. REST API
# ══════════════════════════════════════════════════════════════════════════════
h1('14. REST API')
body(
    'A versioned REST API at /api/v1/ enables integration with HR systems, LMS platforms, '
    'and automation tools. All endpoints require an API key in the Authorization header.'
)

h2('14.1 Core Endpoints')
add_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET',    '/api/v1/exams',                    'List all published exams'],
        ['GET',    '/api/v1/exams/:id',                'Get exam details + sections'],
        ['POST',   '/api/v1/candidates',               'Create candidate record'],
        ['POST',   '/api/v1/links',                    'Generate exam link for candidate'],
        ['GET',    '/api/v1/results',                  'List submissions (filterable)'],
        ['GET',    '/api/v1/results/:id',              'Get full submission detail'],
        ['GET',    '/api/v1/candidates/:id/results',   'All results for one candidate'],
        ['POST',   '/api/v1/webhooks',                 'Register webhook URL for events'],
        ['DELETE', '/api/v1/links/:token',             'Revoke an exam link'],
        ['GET',    '/api/v1/analytics/exams/:id',      'Get exam analytics data'],
    ],
    col_widths=[0.8, 2.8, 3.2]
)

h2('14.2 Webhook Events')
for b in [
    'exam.submitted — fired when candidate completes exam',
    'result.ready — fired when all manual review complete, final score available',
    'link.expired — fired when an unused link passes its expiry',
    'review.requested — fired when text answers need checker review',
]:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 15. GAMIFICATION
# ══════════════════════════════════════════════════════════════════════════════
h1('15. Gamification (Practice Mode)')
add_table(
    ['Feature', 'Description'],
    [
        ['Leaderboard',           'Top scores on practice exams — opt-in per exam, shown on catalog'],
        ['Badges',                'Earned on: first pass, perfect score, speed run, streak milestones'],
        ['Streak tracking',       'Consecutive days with at least one exam activity'],
        ['Progress bars',         '% of questions in bank practiced; per-domain coverage'],
        ['Recommended next exam', 'Based on domain weaknesses from last result, suggests most relevant exam'],
        ['Learning paths',        'Ordered sequences — e.g. AZ-900 → AZ-104 → AZ-305; locks unlock on pass'],
    ],
    col_widths=[2.0, 4.8]
)

# ══════════════════════════════════════════════════════════════════════════════
# 16. DATABASE SCHEMA OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('16. Database Schema Overview')
body('The platform uses 22 SQLite tables:')

add_table(
    ['Table', 'Purpose'],
    [
        ['users',               'Admin, checker, and viewer accounts with roles and 2FA secrets'],
        ['sessions',            'Active JWT sessions for force-logout capability'],
        ['exams',               'All exam configuration — settings, timing, scoring, anti-cheat flags'],
        ['sections',            'Sections/domains within each exam with colour and weight'],
        ['questions',           'All question types with code blocks, marks, difficulty, order'],
        ['question_options',    'MCQ/multi-select/T-F options with is_correct flag'],
        ['question_versions',   'Historical versions of questions for rollback'],
        ['question_bank',       'Global reusable questions not tied to a specific exam'],
        ['candidates',          'Candidate name, email, department, tags'],
        ['departments',         'Organisational groupings for candidates'],
        ['exam_links',          'Unique tokens — one per candidate per exam, single-use'],
        ['submissions',         'One per completed exam attempt — aggregate scores, status'],
        ['answers',             'Individual question responses — marks, review status, remarks'],
        ['review_assignments',  'Which checker is assigned to which submission'],
        ['email_config',        'Azure Graph credentials and default email settings'],
        ['email_templates',     'HTML templates for all 7 email types'],
        ['email_log',           'Log of every email sent with status and timestamp'],
        ['settings',            'Key-value platform configuration'],
        ['api_keys',            'REST API keys with permissions and rate limits'],
        ['webhooks',            'Registered webhook URLs and event subscriptions'],
        ['audit_log',           'Complete action history across all admin users'],
        ['gamification',        'Badges earned, streaks, leaderboard entries per candidate'],
    ],
    col_widths=[2.0, 4.8]
)

# ══════════════════════════════════════════════════════════════════════════════
# 17. BUILD PHASES
# ══════════════════════════════════════════════════════════════════════════════
h1('17. Build Phases')

add_table(
    ['Phase', 'Deliverable', 'Key Components'],
    [
        ['1',  'Foundation',               'Node.js + Express, SQLite schema, auth (login/JWT/2FA/SSO), admin shell + sidebar, audit log'],
        ['2',  'Exam Authoring',           'Exam CRUD, sections, all 8 question types, drag-drop reorder, version history, exam preview, templates'],
        ['3',  'Import & Question Bank',   'CSV import + validation, sample CSV, question bank, auto-generate from bank, question review workflow'],
        ['4',  'Candidates & Links',       'Candidate management (dept/tags/bulk), link generation (pools/password/IP/time), duplicate detection'],
        ['5',  'Exam Taking Engine',       'All 8 question types rendered, anti-cheat suite, multi-monitor detection, save & resume, offline PWA, accessibility'],
        ['6',  'Scoring & Submissions',    'Auto-scoring (all types), negative marking, weighted sections, suspicious activity logging, webcam snapshots'],
        ['7',  'Checker & Review',         'Multi-checker queue, disagreement escalation, grace marks, text/file review interface'],
        ['8',  'Results, PDF & Email',     'Result calculation, jsPDF report, certificate, Azure Graph integration, all 7 email templates, candidate portal'],
        ['9',  'Analytics & Reports',      'Item analysis, heatmaps, cohort compare, custom report builder, scheduled Excel export'],
        ['10', 'White-Label & Branding',   'Custom theme per exam, watermark, certificate builder, custom domain support'],
        ['11', 'API, Security & GDPR',     'REST API, API keys, webhooks, GDPR export, data retention, encrypted answers, session management'],
        ['12', 'Gamification & Extras',    'Leaderboard, badges, streaks, learning paths, post-exam resources, co-author, collaboration'],
        ['13', 'Polish & Documentation',   'Mobile responsive, loading states, onboarding wizard, error handling, inline help'],
    ],
    col_widths=[0.5, 1.8, 4.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 18. PROJECT SCALE
# ══════════════════════════════════════════════════════════════════════════════
h1('18. Project Scale Estimate')
add_table(
    ['Metric', 'Estimate'],
    [
        ['HTML pages',             '~38'],
        ['JavaScript files',       '~28'],
        ['API endpoints',          '~95+'],
        ['Database tables',        '22'],
        ['Email templates',        '7'],
        ['Total source files',     '~130+'],
    ],
    col_widths=[2.5, 4.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# 19. OUT OF SCOPE (v1)
# ══════════════════════════════════════════════════════════════════════════════
h1('19. Out of Scope — Version 1')
body('The following are intentionally excluded from v1 and planned for v2:')
for b in [
    'Microsoft Teams / Slack webhook notifications',
    'SMS notifications (Twilio)',
    'Native mobile application (iOS / Android)',
    'LTI integration for LMS platforms (Moodle, Canvas)',
    'AI-assisted question generation',
    'Video conferencing proctoring (beyond webcam snapshots)',
    'Multi-tenant SaaS architecture (v1 is single-organisation)',
]:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '002B5C')
r = p.add_run(f'  Alaric Exam Platform — Project Specification v1.0 — {datetime.date.today().strftime("%B %Y")}  ')
r.font.size = Pt(9)
r.font.color.rgb = MID_GREY
r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('CONFIDENTIAL — Internal Use Only')
r2.font.size = Pt(9)
r2.font.color.rgb = MID_GREY
r2.font.name = 'Calibri'
r2.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/Users/pete/Downloads/Alaric Exam/Alaric_Exam_Project_Specification_v1.0.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
